import typer
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_BREAK
from PIL import Image
import io
from rich.console import Console

from utils.batch import BatchProcessor
from utils.processor import process_file
from utils.segment_handler import SegmentHandler

console = Console()

# Update the grey color to a lighter shade
GREY_FILL = "E8E8E8"  # Lighter grey

# Initialize docs_dict at module level
docs_dict = {}

def set_document_properties(doc):
    """Set up initial document properties"""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.different_first_page_header_footer = True
    # No default margins - we'll set them per section

def create_cover_page(doc, folder_name):
    """Create a white cover page with centered folder name as title"""
    # Create blank first page
    blank_section = doc.sections[0]
    blank_section.left_margin = blank_section.right_margin = blank_section.top_margin = blank_section.bottom_margin = Inches(1)
    doc.add_paragraph()  # Add empty paragraph for blank page
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    
    # Create new section for cover page
    cover_section = doc.add_section(WD_SECTION.NEW_PAGE)
    cover_section.left_margin = cover_section.right_margin = cover_section.top_margin = cover_section.bottom_margin = Inches(1)
    
    # Add vertical space to center content
    for _ in range(10):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Use folder name as title, replacing underscores with spaces
    title = folder_name.replace('_', ' ')
    run = p.add_run(title)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(24)
    run.font.bold = True
    
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def get_base_filename(filename):
    """Extract base filename without path or extension and remove duplicates"""
    base = Path(filename).stem
    # Split by potential duplicate segments and take the first occurrence
    parts = base.split('_')
    seen = set()
    cleaned_parts = []
    for part in parts:
        if part not in seen:
            seen.add(part)
            cleaned_parts.append(part)
    return '_'.join(cleaned_parts)

def calculate_optimal_font_size(text_length, page_width_inches, page_height_inches):
    """Calculate optimal font size to fit text on one page"""
    # Approximate characters per line and lines per page at 12pt
    CHAR_WIDTH_12PT = 0.11
    LINE_HEIGHT_12PT = 0.18
    
    # Available space (accounting for margins)
    available_width = page_width_inches - 2
    available_height = page_height_inches - 2
    
    # Define font size ranges based on text length
    if text_length < 500:
        font_sizes = list(range(12, 10, -1))    # Short texts
    elif text_length < 1000:
        font_sizes = list(range(11, 9, -1))     # Medium texts
    else:
        font_sizes = list(range(10, 8, -1))     # Long texts
    
    for font_size in font_sizes:
        scale_factor = font_size / 12.0
        chars_per_line = int(available_width / (CHAR_WIDTH_12PT * scale_factor))
        lines_per_page = int(available_height / (LINE_HEIGHT_12PT * scale_factor))
        total_capacity = chars_per_line * lines_per_page
        
        if text_length <= total_capacity:
            return font_size
    
    return 8

def create_spread(doc, image_path, text, filename):
    """Create a two-page spread with image (left) and text (right)"""
    display_name = get_base_filename(filename)
    
    # Left page - Image
    section = doc.add_section()
    
    # Center image by setting equal margins
    margin = Inches(0.75)
    section.left_margin = section.right_margin = margin
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.75)
    
    # Add image paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    
    try:
        with Image.open(image_path) as img:
            # Convert to RGB (removing alpha channel)
            if img.mode in ('RGBA', 'LA') or (img.mode == 'P' and 'transparency' in img.info):
                bg = Image.new('RGB', img.size, 'white')
                if img.mode != 'RGBA':
                    img = img.convert('RGBA')
                bg.paste(img, mask=img.split()[3])
                img = bg
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Calculate dimensions for page
            available_width = 8.5 - (2 * margin.inches)
            available_height = 9.75
            img_ratio = img.height / img.width
            
            if img_ratio > available_height/available_width:
                height = available_height
                width = height / img_ratio
            else:
                width = available_width
                height = width * img_ratio
            
            # Keep 300 DPI resolution for print quality
            target_width = int(width * 300)
            target_height = int(height * 300)
            img = img.resize((target_width, target_height), Image.Resampling.LANCZOS)
            
            # Save as JPG with compression
            img_byte_arr = io.BytesIO()
            img.save(img_byte_arr, format='JPEG', quality=80, optimize=True)
            img_byte_arr.seek(0)
            
            run.add_picture(img_byte_arr, width=Inches(width), height=Inches(height))
    except Exception as e:
        console.print(f"[red]Error loading image {image_path}: {str(e)}")
        run.add_text(f"[Error loading image: {str(e)}]")

    # Right page - Text
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.left_margin = section.right_margin = Inches(1)
    section.top_margin = section.bottom_margin = Inches(1)
    
    # Add filename as title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run(display_name)
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    title_p.space_after = Pt(12)  # Add space after title
    
    # Add text content
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(calculate_optimal_font_size(len(text), 8.5, 11.0))
    
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    
    # Add page break for next spread
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    
    return True

def process_image(file_path: Path, out_path: Path, transcriptions_folder: Path) -> dict:
    """Process a single image file, returning manifest-compatible output"""
    try:
        # Get source folder structure from input path, preserving everything after 'documents'
        source_dir = Path(*file_path.parts[file_path.parts.index('documents'):])
        
        # Find corresponding text file - use the full path from documents/
        text_path = transcriptions_folder / source_dir.with_suffix('.txt')
        
        if not text_path.exists():
            console.print(f"[red]No transcription found for {file_path}. Looked in {text_path}")
            return {
                "error": f"No transcription found for {file_path}. Looked in {text_path}",
                "source": str(file_path)
            }

        # Read transcription text
        transcription_text = text_path.read_text(encoding='utf-8')
        
        # Create or get document
        global docs_dict  # Use the module-level docs_dict
        if str(source_dir.parent) not in docs_dict:
            doc = Document()
            create_cover_page(doc, str(source_dir.parent))
            docs_dict[str(source_dir.parent)] = {"doc": doc}
        
        # Add spread to document
        doc_info = docs_dict[str(source_dir.parent)]
        create_spread(doc_info["doc"], file_path, transcription_text, file_path.stem)
        
        # Build output path preserving full source hierarchy
        rel_path = source_dir.parent.with_suffix('.docx')
        
        return {
            "outputs": [str(rel_path)],
            "source": str(SegmentHandler.get_relative_path(file_path)),
            "details": {
                "text_length": len(transcription_text)
            }
        }
        
    except Exception as e:
        console.print(f"[red]Error processing {file_path}: {e}")
        return {
            "error": str(e),
            "source": str(file_path)
        }

def process_document(file_path: str, output_folder: Path, transcriptions_folder: Path) -> dict:
    """Process a single document file"""
    try:
        # Convert file_path to Path and get base filename
        file_path = Path(file_path)
        base_filename = get_base_filename(file_path.name)
        
        # Get the relative path from the documents folder, preserving everything after 'documents'
        rel_path = Path(*file_path.parts[file_path.parts.index('documents'):])
        
        # Find the corresponding transcription text file
        # Remove .txt extension if it exists and replace with image extension
        if rel_path.suffix == '.txt':
            rel_path = rel_path.with_suffix('')
        
        # Look for image file with various extensions
        image_path = None
        for ext in ['.jpg', '.jpeg', '.png', '.tif', '.tiff']:
            test_path = file_path.with_suffix(ext)
            if test_path.exists():
                image_path = test_path
                break
                
        if not image_path:
            console.print(f"[red]No image found for {file_path}")
            return {
                "error": f"No image found for {file_path}",
                "source": str(file_path)
            }
        
        # Find transcription file - use the full path from documents/
        transcription_file = transcriptions_folder / rel_path.with_suffix('.txt')
        
        if not transcription_file.exists():
            console.print(f"[red]No transcription found for {file_path}. Looked in {transcription_file}")
            return {
                "error": f"No transcription found for {file_path}. Looked in {transcription_file}",
                "source": str(file_path)
            }

        # Read transcription text
        transcription_text = transcription_file.read_text(encoding='utf-8')
        
        # Get the parent folder name for document grouping
        parent_folder = file_path.parent.name
        
        # Create or get document
        if parent_folder not in docs_dict:
            doc = Document()
            set_document_properties(doc)
            create_cover_page(doc, parent_folder)
            docs_dict[parent_folder] = {"doc": doc}
        
        # Add spread to document
        doc_info = docs_dict[parent_folder]
        create_spread(doc_info["doc"], image_path, transcription_text, base_filename)
        
        # Build output path
        out_path = Path(parent_folder).with_suffix('.docx')
        
        return {
            "outputs": [str(out_path)],
            "source": str(file_path),
            "details": {
                "text_length": len(transcription_text)
            }
        }
    except Exception as e:
        console.print(f"[red]Error processing {file_path}: {e}")
        return {
            "error": str(e),
            "source": str(file_path)
        }

def convert_to_word(
    background_removed_folder: Path = typer.Argument(..., help="Input background removed images folder"),
    background_removed_manifest: Path = typer.Argument(..., help="Input background removed manifest"),
    word_folder: Path = typer.Argument(..., help="Output folder for Word documents"),
    transcriptions_folder: Path = typer.Argument(..., help="Folder containing transcription files")
):
    """Convert background-removed images and transcriptions to Word documents with side-by-side layout"""
    console.print(f"[green]Converting images in {background_removed_folder} to Word documents")
    
    # Clear any existing docs_dict
    docs_dict.clear()
    
    processor = BatchProcessor(
        input_manifest=background_removed_manifest,
        output_folder=word_folder,
        process_name="convert_to_word",
        base_folder=background_removed_folder / "documents",
        processor_fn=lambda f, o: process_document(f, o, transcriptions_folder)
    )
    
    results = processor.process()
    
    # Save accumulated documents
    word_folder.mkdir(parents=True, exist_ok=True)
    for doc_key, doc_info in docs_dict.items():
        # Create output path
        out_path = word_folder / Path(doc_key).with_suffix('.docx')
        # Ensure parent directories exist
        out_path.parent.mkdir(parents=True, exist_ok=True)
        # Save document
        doc_info["doc"].save(str(out_path))
    
    return results

if __name__ == "__main__":
    typer.run(convert_to_word)
