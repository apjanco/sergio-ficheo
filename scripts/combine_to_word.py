import srsly
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT, WD_SECTION
from PIL import Image
import typer
import docx.oxml.shared
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = typer.Typer()

def set_page_size(doc, width, height, left_margin, right_margin, top_margin, bottom_margin):
    section = doc.sections[-1]
    section.page_width = Inches(width)
    section.page_height = Inches(height)
    section.orientation = WD_ORIENT.PORTRAIT
    section.left_margin = Inches(left_margin)
    section.right_margin = Inches(right_margin)
    section.top_margin = Inches(top_margin)
    section.bottom_margin = Inches(bottom_margin)

def set_paragraph_format(paragraph, style):
    paragraph.style = style
    paragraph.paragraph_format.space_after = Pt(9)

def format_ner_data(paragraph, ner_data):
    label_map = {
        "LOC": "Location",
        "ORG": "Organization",
        "PER": "Person"
    }
    grouped_entities = {}
    for ent in ner_data:
        label = ent['label']
        if label not in grouped_entities:
            grouped_entities[label] = []
        grouped_entities[label].append(ent['text'])

    for label, texts in grouped_entities.items():
        label_name = label_map.get(label, label)
        run = paragraph.add_run(f"{label_name}:\n")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        text_run = paragraph.add_run("\n".join(texts) + "\n\n")
        text_run.font.name = 'Times New Roman'
        text_run.font.size = Pt(9)

def strip_quotes(text):
    return text.strip().strip('"')

def create_styles(doc):
    styles = doc.styles

    # Heading style
    heading_style = styles.add_style('HeadingStyle', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    heading_font = heading_style.font
    heading_font.name = 'Times New Roman'
    heading_font.size = Pt(14)
    heading_font.bold = True

    # Spanish text style
    spanish_style = styles.add_style('Spanish', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    spanish_font = spanish_style.font
    spanish_font.name = 'Times New Roman'
    spanish_font.size = Pt(9)

    # English text style
    english_style = styles.add_style('English', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    english_font = english_style.font
    english_font.name = 'Times New Roman'
    english_font.size = Pt(9)

    # Summary style
    summary_style = styles.add_style('Summary', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    summary_font = summary_style.font
    summary_font.name = 'Times New Roman'
    summary_font.size = Pt(9)
    summary_font.italic = True

    # NER style
    ner_style = styles.add_style('NER', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    ner_font = ner_style.font
    ner_font.name = 'Times New Roman'
    ner_font.size = Pt(9)

def remove_table_borders(table):
    tbl = table._element
    for cell in tbl.xpath(".//w:tc"):
        tcPr = cell.get_or_add_tcPr()
        tcBorders = docx.oxml.shared.OxmlElement('w:tcBorders')
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = docx.oxml.shared.OxmlElement(f'w:{border_name}')
            border.set(docx.oxml.shared.qn('w:val'), 'nil')
            tcBorders.append(border)
        tcPr.append(tcBorders)

def combine_to_word(json_file: Path, image_folder: Path, output_folder: Path):
    data = list(srsly.read_jsonl(json_file))  # Process all items
    doc_dict = {
        "SM_NPQ_C01_": Document(),
        "SM_NPQ_C02_": Document(),
        "SM_NPQ_C03_": Document(),
        "SM_NPQ_C04_": Document(),
        "SM_NPQ_C05_": Document()
    }

    for doc in doc_dict.values():
        create_styles(doc)
        set_page_size(doc, 10, 11, 0.5, 0.5, 0.5, 0.5)  # Set left-hand page size to 10x11 inches with 0.5 inch margins

    for item in data:
        image_path = image_folder / item.get('image', '')
        if not image_path.exists():
            print(f"Image {image_path} not found. Skipping.")
            continue

        stem_parts = image_path.stem.split('_')
        stem = f"{stem_parts[0]}_{stem_parts[1]}_{stem_parts[2]}_"
        if stem not in doc_dict:
            print(f"Image {image_path} does not match any known stem. Skipping.")
            continue

        doc = doc_dict[stem]

        # Add image to the left-hand page
        image = Image.open(image_path)
        width, height = image.size
        aspect_ratio = width / height
        max_width = Inches(9.75)  # Slightly smaller width to avoid cutting off
        max_height = Inches(10.75)  # Slightly smaller height to avoid cutting off
        if aspect_ratio > max_width / max_height:
            picture = doc.add_picture(str(image_path), width=max_width)
        else:
            picture = doc.add_picture(str(image_path), height=max_height)
        
        # Center the image
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add a new section for the right-hand page
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        set_page_size(doc, 10, 11, 0.5, 0.5, 0.5, 0.5)  # Set right-hand page size to 10x11 inches with 0.5 inch margins

        # Add text boxes to the right-hand page
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'

        # Remove table borders
        remove_table_borders(table)

        # Center the table
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add text
        text_cell = table.cell(0, 0)
        text_cell.text = strip_quotes(item.get('text', 'No text available'))
        set_paragraph_format(text_cell.paragraphs[0], 'Spanish')

        # Add cleaned text
        cleaned_text_cell = table.cell(0, 1)
        cleaned_text_cell.text = strip_quotes(item.get('cleaned_text', 'No cleaned text available'))
        set_paragraph_format(cleaned_text_cell.paragraphs[0], 'Spanish')

        # Add translated text
        translated_text_cell = table.cell(0, 2)
        translated_text_cell.text = strip_quotes(item.get('english_translation', 'No translation available'))
        set_paragraph_format(translated_text_cell.paragraphs[0], 'English')

        # Add NER data
        ner_data_cell = table.cell(1, 0)
        ner_data_paragraph = ner_data_cell.paragraphs[0]
        ner_data_paragraph.style = 'NER'
        ner_data = item.get('entities', 'No NER data available')
        format_ner_data(ner_data_paragraph, ner_data)

        # Add summary
        summary_cell = table.cell(1, 1)
        summary = strip_quotes(item.get('summary', 'No summary available'))
        summary_cell.text = summary
        set_paragraph_format(summary_cell.paragraphs[0], 'Summary')

        # Add file name
        file_name_cell = table.cell(1, 2)
        file_name_cell.text = f"File Name: {image_path.name}"
        set_paragraph_format(file_name_cell.paragraphs[0], 'Summary')

        doc.add_page_break()

    for stem, doc in doc_dict.items():
        output_file = output_folder / f"{stem[:-1]}.docx"
        doc.save(output_file)
        print(f"Document saved as {output_file}")

@app.command()
def main(
    json_file: Path = typer.Argument(..., help="Path to the JSONL file"),
    image_folder: Path = typer.Argument(..., help="Path to the folder containing images"),
    output_folder: Path = typer.Argument(..., help="Path to the output folder for Word files")
):
    combine_to_word(json_file, image_folder, output_folder)

if __name__ == "__main__":
    typer.run(main)