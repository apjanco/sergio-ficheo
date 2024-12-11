import srsly
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT, WD_SECTION
import typer
import docx.oxml.shared
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = typer.Typer()

def set_page_size(doc, width, height, left_margin, right_margin, top_margin, bottom_margin):
    section = doc.sections[-1]
    section.page_width = Inches(width)
    section.page_height = Inches(height)
    section.orientation = WD_ORIENT.PORTRAIT
    section.left_margin = Inches(left_margin)
    section.right_margin = Inches(right_margin)
    section.top_margin = Inches(top_margin)
    section.bottom_margin = Inches(bottom_margin)

def set_paragraph_format(paragraph, style):
    paragraph.style = style
    paragraph.paragraph_format.space_after = Pt(9)

def strip_quotes(text):
    return text.strip().strip('"')

def create_styles(doc):
    styles = doc.styles

    # Heading style
    heading_style = styles.add_style('HeadingStyle', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    heading_font = heading_style.font
    heading_font.name = 'Times New Roman'
    heading_font.size = Pt(14)
    heading_font.bold = True

    # Spanish text style
    spanish_style = styles.add_style('Spanish', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    spanish_font = spanish_style.font
    spanish_font.name = 'Times New Roman'
    spanish_font.size = Pt(9)

    # English text style
    english_style = styles.add_style('English', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    english_font = english_style.font
    english_font.name = 'Times New Roman'
    english_font.size = Pt(9)

    # Summary style
    summary_style = styles.add_style('Summary', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    summary_font = summary_style.font
    summary_font.name = 'Times New Roman'
    summary_font.size = Pt(9)
    summary_font.italic = True

def remove_table_borders(table):
    tbl = table._element
    for cell in tbl.xpath(".//w:tc"):
        tcPr = cell.get_or_add_tcPr()
        tcBorders = docx.oxml.shared.OxmlElement('w:tcBorders')
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = docx.oxml.shared.OxmlElement(f'w:{border_name}')
            border.set(docx.oxml.shared.qn('w:val'), 'nil')
            tcBorders.append(border)
        tcPr.append(tcBorders)

def add_page_number(doc):
    section = doc.sections[-1]
    footer = section.footer.paragraphs[0]
    footer.text = "Page "
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fldChar = docx.oxml.shared.OxmlElement('w:fldChar')
    fldChar.set(docx.oxml.shared.qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = docx.oxml.shared.OxmlElement('w:instrText')
    instrText.text = "PAGE"
    run._r.append(instrText)
    fldChar = docx.oxml.shared.OxmlElement('w:fldChar')
    fldChar.set(docx.oxml.shared.qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def combine_to_word_pretty(json_file: Path, output_folder: Path):
    if not output_folder.exists():
        output_folder.mkdir(parents=True, exist_ok=True)

    data = list(srsly.read_jsonl(json_file))  # Process all items
    doc_dict = {
        "SM_NPQ_C01_": Document(),
        "SM_NPQ_C02_": Document(),
        "SM_NPQ_C03_": Document(),
        "SM_NPQ_C04_": Document(),
        "SM_NPQ_C05_": Document()
    }

    for doc in doc_dict.values():
        create_styles(doc)
        set_page_size(doc, 8.5, 11, 0.5, 0.5, 0.5, 0.5)  # Set page size to 8.5x11 inches with 0.5 inch margins
        add_page_number(doc)

    for item in data:
        stem_parts = item.get('image', '').split('_')
        stem = f"{stem_parts[0]}_{stem_parts[1]}_{stem_parts[2]}_"
        if stem not in doc_dict:
            print(f"Item {item} does not match any known stem. Skipping.")
            continue

        doc = doc_dict[stem]

        # Add a new section for the page
        doc.add_section(WD_SECTION.NEW_PAGE)

        # Add heading with file name
        heading = doc.add_heading(level=1)
        heading.text = f"File Name: {item.get('image', 'No file name available')}"
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.style = 'HeadingStyle'

        # Add text boxes to the page
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'

        # Remove table borders
        remove_table_borders(table)

        # Center the table
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add cleaned text
        cleaned_text_cell = table.cell(0, 0)
        cleaned_text_cell.text = strip_quotes(item.get('cleaned_text', 'No cleaned text available'))
        set_paragraph_format(cleaned_text_cell.paragraphs[0], 'Spanish')

        # Add translated text
        translated_text_cell = table.cell(0, 1)
        translated_text_cell.text = strip_quotes(item.get('english_translation', 'No translation available'))
        set_paragraph_format(translated_text_cell.paragraphs[0], 'English')

        # Add summary
        summary_cell = table.cell(0, 2)
        summary = strip_quotes(item.get('summary', 'No summary available'))
        summary_cell.text = summary
        set_paragraph_format(summary_cell.paragraphs[0], 'Summary')

        # Add people entities
        people_cell = table.cell(0, 3)
        people = ', '.join(ent['text'] for ent in item.get('entities', []) if ent['label'] == 'PER')
        people_cell.text = people
        set_paragraph_format(people_cell.paragraphs[0], 'Summary')

        # Add locations entities
        locations_cell = table.cell(0, 4)
        locations = ', '.join(ent['text'] for ent in item.get('entities', []) if ent['label'] == 'LOC')
        locations_cell.text = locations
        set_paragraph_format(locations_cell.paragraphs[0], 'Summary')

        # Add organizations entities
        orgs_cell = table.cell(0, 5)
        orgs = ', '.join(ent['text'] for ent in item.get('entities', []) if ent['label'] == 'ORG')
        orgs_cell.text = orgs
        set_paragraph_format(orgs_cell.paragraphs[0], 'Summary')

    for stem, doc in doc_dict.items():
        output_file = output_folder / f"{stem[:-1]}.docx"
        if output_file.exists():
            output_file.unlink()  # Delete the existing file if it exists
        doc.save(output_file)
        print(f"Document saved as {output_file}")

@app.command()
def main(
    json_file: Path = typer.Argument(..., help="Path to the JSONL file"),
    output_folder: Path = typer.Argument(..., help="Path to the output folder for Word files")
):
    combine_to_word_pretty(json_file, output_folder)

if __name__ == "__main__":
    typer.run(main)